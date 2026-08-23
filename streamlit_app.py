
import streamlit as st
import json, re, hashlib, secrets, string, mimetypes
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt, Inches
from supabase import create_client

BASE=Path(__file__).parent
with open(BASE/"modulos.json",encoding="utf-8") as f: MODULES=json.load(f)
BUCKET="project-evidence"; MAX_FILE_MB=10
ALLOWED_MIME={"image/png","image/jpeg","application/pdf"}

st.set_page_config(page_title="Project Builder | Desarrollo de Emprendedores",page_icon="🚀",layout="wide",initial_sidebar_state="expanded")
st.markdown("""<style>
.block-container{padding-top:1rem;max-width:1500px}.hero{padding:1.3rem 1.5rem;border-radius:18px;background:#f7f8fa;border:1px solid #e4e6e9;margin-bottom:1rem}
.hero h1{margin:0 0 .2rem 0}.muted{color:#666}.role{display:inline-block;padding:.2rem .55rem;border-radius:14px;background:#eceff2;margin-right:.3rem;font-size:.8rem}
.review-box{padding:.8rem 1rem;border:1px solid #e3e5e8;border-radius:12px;background:#fbfbfc;margin:.5rem 0}
</style>""",unsafe_allow_html=True)

def utc_now(): return datetime.now(timezone.utc).isoformat()
def get_db():
    try: return create_client(st.secrets["SUPABASE_URL"],st.secrets["SUPABASE_SERVICE_KEY"])
    except Exception: st.error("La aplicación no está conectada con Supabase."); st.stop()
db=get_db()

def hash_secret(v,s): return hashlib.pbkdf2_hmac("sha256",v.encode("utf-8"),bytes.fromhex(s),180000).hex()
def verify_secret(v,s,e):
    try: return secrets.compare_digest(hash_secret(v,s),e)
    except Exception: return False
def new_salt(): return secrets.token_hex(16)
def random_password(n=10):
    alphabet=string.ascii_letters+string.digits+"!@#$"
    return "".join(secrets.choice(alphabet) for _ in range(n))
def normalize_code(v): return re.sub(r"[^A-Z0-9_-]","",v.strip().upper())
def safe_name(v): return re.sub(r"[^A-Za-z0-9._-]+","_",v).strip("_") or "archivo"

def rows(table,**equals):
    q=db.table(table).select("*")
    for k,v in equals.items(): q=q.eq(k,v)
    return q.execute().data or []
def one(table,**equals):
    r=rows(table,**equals); return r[0] if r else None

def ensure_initial_admin():
    if rows("app_users"): return
    try: email=st.secrets["INITIAL_ADMIN_EMAIL"].strip().lower(); password=st.secrets["INITIAL_ADMIN_PASSWORD"]; name=st.secrets.get("INITIAL_ADMIN_NAME","Administradora")
    except Exception: return
    salt=new_salt()
    db.table("app_users").insert({"name":name,"email":email,"password_salt":salt,"password_hash":hash_secret(password,salt),"is_admin":True,"is_teacher":True,"is_active":True,"created_at":utc_now(),"updated_at":utc_now()}).execute()
ensure_initial_admin()

def staff_login(email,password):
    u=one("app_users",email=email.strip().lower(),is_active=True)
    return u if u and verify_secret(password,u["password_salt"],u["password_hash"]) else None
def team_login(code,pin):
    t=one("teams",team_code=normalize_code(code),is_active=True)
    return t if t and verify_secret(pin,t["pin_salt"],t["pin_hash"]) else None
def all_fields(m): return [f for s in m["sections"] for f in s["fields"]]
def word_count(text): return len(re.findall(r"\b[\wÁÉÍÓÚÜÑáéíóúüñ'-]+\b",text or ""))
def module_pct(m,payload):
    md=payload.get("modulos",{}).get(m["id"],{}); fs=all_fields(m)
    return round(100*sum(bool(str(md.get(f["clave"],"")).strip()) for f in fs)/len(fs)) if fs else 0
def global_pct(payload): return round(sum(module_pct(m,payload) for m in MODULES)/len(MODULES)) if MODULES else 0
def consistency_alerts(payload):
    mods=payload.get("modulos",{}); fin=mods.get("finanzas",{}); a=[]
    if mods.get("mercadotecnia",{}).get("presupuesto_mkt","").strip() and not fin.get("presupuesto","").strip(): a.append("Mercadotecnia ya tiene presupuesto; debe integrarse en Finanzas.")
    if mods.get("operaciones",{}).get("proveedores","").strip() and not fin.get("presupuesto","").strip(): a.append("Operaciones ya identifica proveedores/costos; deben integrarse en Finanzas.")
    if mods.get("talento",{}).get("compensacion","").strip() and not fin.get("presupuesto","").strip(): a.append("Talento ya define compensaciones; nómina, prestaciones e incentivos deben integrarse en Finanzas.")
    if mods.get("legal",{}).get("costos_legales","").strip() and not fin.get("presupuesto","").strip(): a.append("Legal ya identifica costos; deben integrarse en Finanzas.")
    return a

# Reviews
def get_review(pid,mid): return one("module_reviews",project_id=pid,module_id=mid)
def save_review(pid,mid,teacher_id,status,feedback):
    old=get_review(pid,mid); data={"status":status,"feedback":feedback.strip() if feedback else None,"reviewed_by":teacher_id,"reviewed_at":utc_now(),"updated_at":utc_now()}
    if old: db.table("module_reviews").update(data).eq("id",old["id"]).execute()
    else:
        data.update({"project_id":pid,"module_id":mid,"created_at":utc_now()})
        db.table("module_reviews").insert(data).execute()
def status_icon(s): return {"En desarrollo":"🟡","Requiere ajustes":"🟠","Revisado":"🟢"}.get(s,"🟡")

# Evidence / Storage
def list_evidences(pid,mid=None,sid=None):
    q=db.table("evidences").select("*").eq("project_id",pid)
    if mid is not None: q=q.eq("module_id",mid)
    if sid is not None: q=q.eq("slot_id",sid)
    return q.execute().data or []
def add_link_evidence(pid,mid,slot,url):
    p=urlparse(url.strip())
    if p.scheme not in ("http","https") or not p.netloc: raise ValueError("La liga debe comenzar con http:// o https://")
    db.table("evidences").insert({"project_id":pid,"module_id":mid,"slot_id":slot["id"],"slot_label":slot["label"],"kind":"link","external_url":url.strip(),"uploaded_at":utc_now()}).execute()
def upload_file_evidence(pid,mid,slot,up):
    mime=up.type or mimetypes.guess_type(up.name)[0] or "application/octet-stream"; raw=up.getvalue()
    if mime not in ALLOWED_MIME: raise ValueError("Solo se permiten PNG, JPG/JPEG o PDF.")
    if len(raw)>MAX_FILE_MB*1024*1024: raise ValueError(f"El archivo excede {MAX_FILE_MB} MB.")
    ext=Path(up.name).suffix.lower()
    path=f"{pid}/{mid}/{slot['id']}/{int(datetime.now().timestamp())}_{secrets.token_hex(4)}_{safe_name(Path(up.name).stem)}{ext}"
    db.storage.from_(BUCKET).upload(path=path,file=raw,file_options={"content-type":mime,"upsert":"false"})
    db.table("evidences").insert({"project_id":pid,"module_id":mid,"slot_id":slot["id"],"slot_label":slot["label"],"kind":"file","file_name":up.name,"storage_path":path,"mime_type":mime,"uploaded_at":utc_now()}).execute()
def delete_evidence(ev):
    if ev.get("kind")=="file" and ev.get("storage_path"):
        try: db.storage.from_(BUCKET).remove([ev["storage_path"]])
        except Exception: pass
    db.table("evidences").delete().eq("id",ev["id"]).execute()
def signed_url(path,expires=3600):
    try:
        r=db.storage.from_(BUCKET).create_signed_url(path,expires)
        return (r.get("signedURL") or r.get("signedUrl") or r.get("signed_url")) if isinstance(r,dict) else None
    except Exception: return None
def download_storage(path): return db.storage.from_(BUCKET).download(path)

def render_evidence_list(pid,mid,allow_delete=False,key_prefix="ev"):
    evs=list_evidences(pid,mid)
    if not evs: st.caption("Aún no hay evidencias cargadas en este módulo."); return
    for ev in evs:
        c1,c2=st.columns([5,1])
        if ev["kind"]=="link": c1.markdown(f"🔗 **{ev['slot_label']}** — {ev.get('external_url','')}")
        else:
            u=signed_url(ev.get("storage_path")); label=ev.get("file_name") or ev["slot_label"]
           if u:
               c1.markdown(f"📎 **{ev['slot_label']}** — [{label}]({u})")
           else:
               c1.write(f"📎 **{ev['slot_label']}** — {label}")
        if allow_delete and c2.button("Eliminar",key=f"{key_prefix}_{ev['id']}"): delete_evidence(ev); st.rerun()

def render_evidence_uploader(pid,m):
    if not m.get("evidence_slots"): return
    st.subheader("📎 Evidencias y anexos"); st.info(m["evidence_note"])
    for slot in m["evidence_slots"]:
        with st.expander(f"➕ {slot['label']}",expanded=False):
            st.caption(slot["description"]); allowed=slot["types"]
            if "image" in allowed or "pdf" in allowed:
                accepted=(["png","jpg","jpeg"] if "image" in allowed else [])+(["pdf"] if "pdf" in allowed else [])
                up=st.file_uploader(f"Subir archivo ({', '.join(accepted).upper()})",type=accepted,key=f"up_{pid}_{m['id']}_{slot['id']}")
                if st.button("Guardar archivo",key=f"upbtn_{pid}_{m['id']}_{slot['id']}"):
                    if not up: st.warning("Selecciona un archivo primero.")
                    else:
                        try: upload_file_evidence(pid,m["id"],slot,up); st.success("Evidencia guardada."); st.rerun()
                        except Exception as e: st.error(f"No fue posible guardar el archivo: {e}")
            if "link" in allowed:
                link=st.text_input("Liga externa (Canva, Drive u otra herramienta)",key=f"lnk_{pid}_{m['id']}_{slot['id']}")
                if st.button("Guardar liga",key=f"lnkbtn_{pid}_{m['id']}_{slot['id']}"):
                    try: add_link_evidence(pid,m["id"],slot,link); st.success("Liga guardada."); st.rerun()
                    except Exception as e: st.error(str(e))
            st.markdown("**Evidencias registradas en este apartado**")
            evs=list_evidences(pid,m["id"],slot["id"])
            if not evs: st.caption("Sin evidencias todavía.")
            for ev in evs:
                if ev["kind"]=="link": st.markdown(f"🔗 {ev.get('external_url','')}")
                else:
                    u = signed_url(ev.get("storage_path"))
                    if u:
                        st.markdown(f"📎 [{ev.get('file_name','Archivo')}]({u})")
                    else:
                        st.write(f"📎 {ev.get('file_name','Archivo')}")
                if st.button("Eliminar evidencia",key=f"del_{ev['id']}"): delete_evidence(ev); st.rerun()

# DOCX
def add_evidences_to_doc(doc,pid,m):
    evs=list_evidences(pid,m["id"])
    if not evs: return
    embed={x["id"]:x.get("embed_images",False) for x in m.get("evidence_slots",[])}
    doc.add_heading("Evidencias y anexos del módulo",2)
    for ev in evs:
        doc.add_heading(ev["slot_label"],3)
        if ev["kind"]=="link": doc.add_paragraph(f"Liga externa: {ev.get('external_url','')}"); continue
        mime=ev.get("mime_type") or ""; filename=ev.get("file_name") or "Archivo"
        if mime.startswith("image/") and embed.get(ev["slot_id"],False):
            try:
                raw=download_storage(ev["storage_path"]); doc.add_picture(BytesIO(raw),width=Inches(5.8)); doc.add_paragraph(f"Evidencia: {filename}")
            except Exception: doc.add_paragraph(f"Imagen registrada: {filename} (no fue posible incrustarla).")
        else: doc.add_paragraph(f"Anexo / evidencia registrada: {filename}")
def build_docx(payload,team,project,pid):
    doc=Document(); doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(11)
    doc.add_heading("Proyecto Final – Plan de Negocios",0); doc.add_paragraph("Desarrollo de Emprendedores"); doc.add_paragraph(f"Equipo: {team}\nProyecto: {project}")
    for m in MODULES:
        md=payload.get("modulos",{}).get(m["id"],{}); doc.add_page_break(); doc.add_heading(f"{m['numero']} · {m['titulo']}",1); doc.add_paragraph(m["objetivo"])
        for s in m["sections"]:
            doc.add_heading(s["titulo"],2)
            for fld in s["fields"]: doc.add_heading(fld["etiqueta"],3); doc.add_paragraph(str(md.get(fld["clave"],"")).strip() or "[Pendiente]")
        add_evidences_to_doc(doc,pid,m); doc.add_heading("Lista de verificación",2); ch=md.get("_checks",{})
        for c in m["checklist"]: doc.add_paragraph(("✓ " if ch.get(c) else "☐ ")+c)
    bio=BytesIO(); doc.save(bio); bio.seek(0); return bio

def save_project(pid,payload): db.table("projects").update({"payload":payload,"updated_at":utc_now()}).eq("id",pid).execute()
def logout():
    for k in list(st.session_state.keys()):
        if k.startswith("auth_") or k in ("current_view","generated_docx"): st.session_state.pop(k,None)
    st.rerun()

# Admin helpers
def create_teacher(name,email):
    email=email.strip().lower()
    if one("app_users",email=email): return None,"Ese correo ya existe."
    temp=random_password(); salt=new_salt()
    db.table("app_users").insert({"name":name.strip(),"email":email,"password_salt":salt,"password_hash":hash_secret(temp,salt),"is_admin":False,"is_teacher":True,"is_active":True,"created_at":utc_now(),"updated_at":utc_now()}).execute()
    return temp,None
def reset_staff_password(uid):
    temp=random_password(); salt=new_salt(); db.table("app_users").update({"password_salt":salt,"password_hash":hash_secret(temp,salt),"updated_at":utc_now()}).eq("id",uid).execute(); return temp
def create_group(name,tid,period): return db.table("groups").insert({"name":name.strip(),"teacher_id":tid,"period":period.strip(),"is_active":True,"created_at":utc_now()}).execute()
def create_team(gid,code,tname,pname,pin):
    code=normalize_code(code)
    if one("teams",team_code=code): return "Ese código de equipo ya existe."
    salt=new_salt(); r=db.table("teams").insert({"group_id":gid,"team_code":code,"team_name":tname.strip(),"pin_salt":salt,"pin_hash":hash_secret(pin,salt),"is_active":True,"created_at":utc_now(),"updated_at":utc_now()}).execute(); tm=r.data[0]
    db.table("projects").insert({"team_id":tm["id"],"project_name":pname.strip(),"payload":{"modulos":{}},"created_at":utc_now(),"updated_at":utc_now()}).execute(); return None
def reset_team_pin(tid):
    pin="".join(secrets.choice(string.digits) for _ in range(6)); salt=new_salt(); db.table("teams").update({"pin_salt":salt,"pin_hash":hash_secret(pin,salt),"updated_at":utc_now()}).eq("id",tid).execute(); return pin

def render_teacher_project(p,user):
    payload=p.get("payload") or {"modulos":{}}; st.write(f"**Proyecto:** {p['project_name']}"); st.write(f"**Avance global:** {global_pct(payload)}%"); st.progress(global_pct(payload)/100)
    for m in MODULES:
        rev=get_review(p["id"],m["id"]); status=rev["status"] if rev else "En desarrollo"
        with st.expander(f"{status_icon(status)} {m['numero']} · {m['titulo']} — {module_pct(m,payload)}%",expanded=False):
            md=payload.get("modulos",{}).get(m["id"],{})
            for s in m["sections"]:
                st.markdown(f"#### {s['titulo']}")
                for f in s["fields"]:
                    value=str(md.get(f["clave"],"")).strip()
                    if value: st.markdown(f"**{f['etiqueta']}**"); st.write(value)
                    else: st.caption(f"{f['etiqueta']}: sin información capturada")
            st.markdown("#### 📎 Evidencias"); render_evidence_list(p["id"],m["id"],False,f"teach_{p['id']}_{m['id']}")
            st.markdown("#### 💬 Retroalimentación docente")
            opts=["En desarrollo","Requiere ajustes","Revisado"]
            with st.form(f"rev_{p['id']}_{m['id']}"):
                ns=st.selectbox("Estatus del módulo",opts,index=opts.index(status) if status in opts else 0)
                nf=st.text_area("Retroalimentación para el equipo",value=rev.get("feedback","") if rev else "",height=130)
                sv=st.form_submit_button("Guardar retroalimentación")
            if sv: save_review(p["id"],m["id"],user["id"],ns,nf); st.success("Retroalimentación guardada."); st.rerun()
            if rev and rev.get("reviewed_at"): st.caption(f"Última revisión docente: {rev['reviewed_at']}")

# LOGIN
if not st.session_state.get("auth_type"):
    st.markdown('<div class="hero"><h1>🚀 Project Builder</h1><b>Desarrollo de Emprendedores · Plan de Negocios</b><div class="muted">Una sola plataforma para profesores, grupos y equipos.</div></div>',unsafe_allow_html=True)
    a,b=st.tabs(["Profesor / Administrador","Equipo"])
    with a:
        email=st.text_input("Correo"); pw=st.text_input("Contraseña",type="password")
        if st.button("Entrar como docente",type="primary",use_container_width=True):
            u=staff_login(email,pw)
            if not u: st.error("Correo o contraseña incorrectos.")
            else: st.session_state.auth_type="staff"; st.session_state.auth_user=u; st.session_state.current_view="Mis grupos"; st.rerun()
    with b:
        code=st.text_input("Código de equipo",key="team_code"); pin=st.text_input("PIN",type="password",key="team_pin")
        if st.button("Entrar al proyecto",use_container_width=True):
            tm=team_login(code,pin)
            if not tm: st.error("Código o PIN incorrecto.")
            else:
                p=one("projects",team_id=tm["id"])
                if not p: st.error("El equipo todavía no tiene proyecto asignado.")
                else: st.session_state.auth_type="team"; st.session_state.auth_team=tm; st.session_state.auth_project=p; st.rerun()
    st.stop()

# STAFF
if st.session_state.auth_type=="staff":
    user=st.session_state.auth_user
    st.markdown(f'<div class="hero"><h1>🚀 Project Builder</h1><b>{user["name"]}</b><br><span class="role">{"Administradora" if user["is_admin"] else ""}</span><span class="role">{"Profesora" if user["is_teacher"] else ""}</span></div>',unsafe_allow_html=True)
    with st.sidebar:
        opts=[] 
        if user["is_teacher"]: opts.append("Mis grupos")
        if user["is_admin"]: opts.append("Panel administrador")
        cur=st.session_state.get("current_view",opts[0]); view=st.radio("Vista",opts,index=opts.index(cur) if cur in opts else 0); st.session_state.current_view=view
        if st.button("Cerrar sesión",use_container_width=True): logout()
    if view=="Panel administrador":
        teachers=[u for u in rows("app_users") if u.get("is_teacher")]; groups=rows("groups"); teams=rows("teams"); projects=rows("projects")
        st.header("Panel Administrador"); c1,c2,c3,c4=st.columns(4); c1.metric("Profesores",len(teachers)); c2.metric("Grupos",len(groups)); c3.metric("Equipos",len(teams)); c4.metric("Proyectos",len(projects))
        A,B,C=st.tabs(["Profesores","Grupos","Seguimiento general"])
        with A:
            with st.form("new_teacher"): name=st.text_input("Nombre"); email=st.text_input("Correo"); sub=st.form_submit_button("Crear profesor")
            if sub:
                if not name.strip() or not email.strip(): st.error("Completa nombre y correo.")
                else:
                    temp,err=create_teacher(name,email)
                    if err: st.error(err)
                    else: st.success(f"Profesor creado. Contraseña temporal: {temp}")
            for t in teachers:
                with st.expander(f"{t['name']} · {t['email']}"):
                    st.write(f"Administradora: {'Sí' if t['is_admin'] else 'No'} · Profesora: {'Sí' if t['is_teacher'] else 'No'}")
                    if t["id"]!=user["id"] and st.button("Restablecer contraseña",key=f"rst_{t['id']}"): st.success(f"Nueva contraseña temporal: {reset_staff_password(t['id'])}")
        with B:
            choices={f"{t['name']} · {t['email']}":t["id"] for t in teachers}
            if choices:
                with st.form("new_group"): gn=st.text_input("Nombre del grupo"); per=st.text_input("Periodo",placeholder="2026-2"); lab=st.selectbox("Profesor responsable",list(choices)); gs=st.form_submit_button("Crear grupo")
                if gs:
                    if not gn.strip(): st.error("Escribe el nombre del grupo.")
                    else: create_group(gn,choices[lab],per); st.success("Grupo creado."); st.rerun()
            for g in rows("groups"):
                t=one("app_users",id=g["teacher_id"]); st.write(f"**{g['name']}** · {g.get('period','')} · {t['name'] if t else 'Sin profesor'}")
        with C:
            gb={g["id"]:g for g in groups}; tb={t["id"]:t for t in teachers}; tmmap={t["id"]:t for t in teams}
            for p in projects:
                tm=tmmap.get(p["team_id"]); g=gb.get(tm["group_id"]) if tm else None; tr=tb.get(g["teacher_id"]) if g else None; v=global_pct(p.get("payload") or {"modulos":{}})
                st.write(f"**{p['project_name']}** · {tm['team_code'] if tm else ''} · {g['name'] if g else ''} · {tr['name'] if tr else ''} · **{v}%**"); st.progress(v/100)
        st.stop()
    st.header("Mis grupos")
    gs=rows("groups",teacher_id=user["id"])
    if not gs: st.info("Aún no tienes grupos asignados.")
    for g in gs:
        with st.expander(f"{g['name']} · {g.get('period','')}",expanded=True):
            with st.form(f"team_{g['id']}"):
                x,y=st.columns(2); code=x.text_input("Código único",key=f"code_{g['id']}"); tn=y.text_input("Nombre del equipo",key=f"tn_{g['id']}"); pn=st.text_input("Nombre del proyecto",key=f"pn_{g['id']}"); pin=st.text_input("PIN inicial",key=f"pin_{g['id']}"); sub=st.form_submit_button("Crear equipo y proyecto")
            if sub:
                if not code.strip() or not tn.strip() or not pn.strip() or len(pin)<4: st.error("Completa todos los datos y usa PIN de al menos 4 caracteres.")
                else:
                    err=create_team(g["id"],code,tn,pn,pin)
                    if err: st.error(err)
                    else: st.success("Equipo creado."); st.rerun()
            for tm in rows("teams",group_id=g["id"]):
                p=one("projects",team_id=tm["id"]); payload=(p or {}).get("payload") or {"modulos":{}}; v=global_pct(payload); c1,c2,c3=st.columns([3,2,1])
                c1.write(f"**{tm['team_code']} · {tm['team_name']}**"); c1.caption(p["project_name"] if p else "Sin proyecto"); c2.progress(v/100); c2.caption(f"{v}% · {(p or {}).get('updated_at','')}")
                if c3.button("Nuevo PIN",key=f"pinreset_{tm['id']}"): st.success(f"Nuevo PIN: {reset_team_pin(tm['id'])}")
                if p:
                    with st.expander("👁️ Ver proyecto y retroalimentar",expanded=False): render_teacher_project(p,user)
    st.stop()

# TEAM
team=st.session_state.auth_team; project=one("projects",id=st.session_state.auth_project["id"]) or st.session_state.auth_project; st.session_state.auth_project=project
payload=project.get("payload") or {"modulos":{}}; payload.setdefault("modulos",{})
group=one("groups",id=team["group_id"]); teacher=one("app_users",id=group["teacher_id"]) if group else None
st.markdown(f'<div class="hero"><h1>🚀 Project Builder</h1><b>{project["project_name"]}</b><br><span class="muted">{team["team_code"]} · {team["team_name"]} · {group["name"] if group else ""} · {teacher["name"] if teacher else ""}</span></div>',unsafe_allow_html=True)
with st.sidebar:
    overall=global_pct(payload); st.metric("Avance global",f"{overall}%"); st.progress(overall/100)
    labels=[]
    for m in MODULES:
        rev=get_review(project["id"],m["id"]); s=rev["status"] if rev else "En desarrollo"; labels.append(f"{m['numero']} · {m['titulo']} · {module_pct(m,payload)}% {status_icon(s)}")
    selected=st.radio("Ruta del proyecto",labels); module=MODULES[labels.index(selected)]
    if st.button("💾 Guardar avance",use_container_width=True): save_project(project["id"],payload); st.success("Guardado en línea.")
    if st.button("📄 Preparar Plan de Negocios",use_container_width=True):
        try: st.session_state.generated_docx=build_docx(payload,team["team_code"],project["project_name"],project["id"]).getvalue(); st.success("Documento preparado.")
        except Exception as e: st.error(f"No fue posible generar el documento: {e}")
    if st.session_state.get("generated_docx"): st.download_button("⬇️ Descargar Plan de Negocios",st.session_state.generated_docx,file_name=f"{team['team_code']}_Plan_de_Negocios.docx",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
    if st.button("Cerrar sesión",use_container_width=True): logout()

t1,t2,t3=st.tabs(["🧭 Construir","🔗 Coherencia","📋 Vista de entrega"])
with t1:
    rev=get_review(project["id"],module["id"]); status=rev["status"] if rev else "En desarrollo"
    st.header(f"{module['numero']} · {module['titulo']}"); st.markdown(f"**Estatus docente:** {status_icon(status)} {status}")
    if rev and (rev.get("feedback") or rev.get("reviewed_at")):
        st.markdown("### 💬 Retroalimentación de tu profesor")
        if rev.get("feedback"): st.write(rev["feedback"])
        if rev.get("reviewed_at"): st.caption(f"Última revisión: {rev['reviewed_at']}")
    st.info("**Objetivo:** "+module["objetivo"])
    with st.expander("♻️ Insumos que debes recuperar",expanded=True):
        for x in module["recupera"]: st.markdown(f"- {x}")
    md=payload["modulos"].setdefault(module["id"],{})
    for s in module["sections"]:
        st.subheader(s["titulo"])
        with st.expander("Indicaciones",expanded=True):
            for x in s["instructions"]: st.markdown(f"- {x}")
        for f in s["fields"]:
            val=st.text_area(f["etiqueta"],md.get(f["clave"],""),help=f.get("ayuda",""),height=145,key=f"{module['id']}__{f['clave']}"); md[f["clave"]]=val
            st.caption(f"Extensión sugerida: {f.get('words_min',120)}–{f.get('words_max',300)} palabras · Actualmente: {word_count(val)} palabras. La extensión es orientativa, no bloqueante.")
    render_evidence_uploader(project["id"],module)
    st.subheader("🤖 Uso responsable de IA"); a,b=st.columns(2)
    with a:
        st.markdown("**Sí puede apoyar en:**")
        for x in module["ia_si"]: st.markdown(f"- {x}")
    with b:
        st.markdown("**No debe utilizarse para:**")
        for x in module["ia_no"]: st.markdown(f"- {x}")
    with st.expander("🤖 Revisar mi avance con IA",expanded=False):
        st.info("La IA revisa y cuestiona; no sustituye el trabajo del equipo.")
        content="\n".join(f"{f['etiqueta']}: {md.get(f['clave'],'')}" for f in all_fields(module))
        prompt=f"""Actúa como asesor académico del curso Desarrollo de Emprendedores.
Revisa el módulo "{module['titulo']}" del proyecto "{project['project_name']}".
No escribas el entregable por el equipo ni inventes datos, fuentes, cifras, competidores, proveedores o requisitos.

1. Identifica inconsistencias.
2. Señala evidencia faltante.
3. Revisa conexión con módulos previos.
4. Distingue dato comprobado, supuesto e hipótesis.
5. Formula cinco preguntas concretas para mejorar.

Contenido:
{content}"""
        st.code(prompt,language=None)
    st.subheader("✅ Checklist"); ch=md.setdefault("_checks",{})
    for c in module["checklist"]: ch[c]=st.checkbox(c,ch.get(c,False),key=f"{module['id']}__c__{c}")
    st.caption("El checklist ayuda a revisar el cumplimiento; no es obligatorio marcarlo para guardar.")
    if st.button("Guardar este módulo",type="primary"): save_project(project["id"],payload); st.session_state.pop("generated_docx",None); st.success("Módulo guardado.")
with t2:
    st.header("Coherencia transversal"); alerts=consistency_alerts(payload)
    if alerts:
        for a in alerts: st.warning(a)
    else: st.success("No hay alertas básicas pendientes.")
    for m in MODULES:
        v=module_pct(m,payload); r=get_review(project["id"],m["id"]); s=r["status"] if r else "En desarrollo"; st.write(f"**{m['numero']} · {m['titulo']}** — {v}% · {status_icon(s)} {s}"); st.progress(v/100)
with t3:
    st.header("Vista de entrega")
    for m in MODULES:
        v=module_pct(m,payload); r=get_review(project["id"],m["id"]); s=r["status"] if r else "En desarrollo"
        with st.expander(f"{m['numero']} · {m['titulo']} — {v}% · {status_icon(s)} {s}"):
            ch=payload.get("modulos",{}).get(m["id"],{}).get("_checks",{})
            for c in m["checklist"]: st.write(("✅ " if ch.get(c) else "⬜ ")+c)
            st.markdown("**Evidencias registradas**"); render_evidence_list(project["id"],m["id"],False,f"delivery_{m['id']}")
